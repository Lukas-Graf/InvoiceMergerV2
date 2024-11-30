""" 
This Module holds all methods to create an invoice

File is written in pylint standard
"""

import os
import time
import random
import shutil
import subprocess
from datetime import date

from docx import Document
from docx2pdf import convert
from docx.shared import Inches

import logger as log
from Config import Config



class Invoice(Config):
    """
    Class defining methods to create an invoice
    
    ...

    Methods
    -------
    write_image:
        Writes all detected tables onto the invoice
    write_text:
        Writes all the needed text (cost of parts, etc.) on the invoice
    pipeline:
        Full pipeline that executes the whole invoice-creation

    Private Methods
    ---------------
    __fill_variables:
        Fills variables like email, invoice_nr, etc. on the invoice
    """

    def __init__(self) -> None:
        super().__init__()
        self.__logger = log.get_logger()


    def write_image(self, doc: Document) -> None:
        """
        Writes all detected tables onto the invoice
        
        Parameters
        ----------
        doc : Document (docx module)
            Document where images should be written on

        Returns
        -------
        None
        """

        page_width = doc.sections[0].page_width.inches
        left_margin = doc.sections[0].left_margin.inches
        right_margin = doc.sections[0].right_margin.inches
        max_width = page_width - left_margin - right_margin

        for picture in os.listdir(f"{self.folder_src()}/temp/"):
            if str(picture).startswith("table"):
                doc.add_picture(f"{self.folder_src()}/temp/{picture}", width=Inches(max_width))
                self.__logger.info(f"Table '{picture}' was written on the invoice")
                
        return None


    def write_text(self, doc: Document, email: str, part_cost: float,
                   personal_cost: float, cost_deduction: float) -> None:
        """
        Writes all the needed text (cost of parts, etc.) on the invoice
        
        Parameters
        ----------
        doc : Document (docx module)
            Document where images should be written on
        email : str
            Email adress for paypal
        part_cost : float
            Total sum of the parts
        personal_cost : float
            Total sum of personal loan
        cost_deduction : float (default=0.00)
            Deduction of single costs, e.g. Oil was bought but
            only 5 from 10 liters are used -> The price from the
            other 5 liters is substracted from the total

        Returns
        -------
        None
        """

        doc.add_paragraph()

        p1 = doc.add_paragraph()
        p1.paragraph_format.tab_stops.add_tab_stop(Inches(1)) 
        p1.add_run("\tKosten Teile:").bold = True
        p1.add_run(f" {part_cost} €")

        p2 = doc.add_paragraph()
        p2.paragraph_format.tab_stops.add_tab_stop(Inches(1))
        p2.add_run("\tKosten Lohn:").bold = True
        p2.add_run(f" {personal_cost} €")

        p3 = doc.add_paragraph()
        p3.paragraph_format.tab_stops.add_tab_stop(Inches(1))
        p3.add_run("\tAbzug Einzelkosten:").bold = True
        p3.add_run(f" {cost_deduction} €")

        p4 = doc.add_paragraph()
        p4.paragraph_format.tab_stops.add_tab_stop(Inches(1))
        p4.add_run("\tGesamtkosten:").bold = True
        p4.add_run(f" {round(part_cost+personal_cost-cost_deduction, 2)} €")

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Book Antiqua"
                run.font.size = Inches(0.15)

        self.__fill_variables(doc=doc, email=email)
        doc.save("./invoice_final.docx")

        try:
            convert("./invoice_final.docx", "./invoice_final.pdf")
        except NotImplementedError:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "./invoice_final.docx", "--outdir", "./"], check=True)

        return None
    

    def pipeline(self, email: str, part_cost: float =None, hourly_rate: float =20.00,
                 hours_worked: float =0.00, cost_deduction: float =0.00) -> None:
        """
        Full pipeline that executes the whole invoice-creation
        
        Parameters
        ----------
        email : str
            Email address for paypal
        part_cost : float (default=None)
            Total sum of the parts
        hourly_rate : float (default=20.00)
            Hourly rate which the employee charges
        hours_worked : float (default=0.00)
            Hours that the employee worked
        cost_deduction : float (default=0.00)
            Deduction of single costs, e.g. Oil was bought but
            only 5 from 10 liters are used -> The price from the
            other 5 liters is substracted from the total

        Returns
        -------
        None
        """

        start_time = time.time()
        shutil.copy("./Invoice_Template.docx", "./invoice_final.docx")
        doc = Document("./invoice_final.docx")

        self.write_image(doc=doc)
        self.write_text(doc=doc,
                        email=email,
                        part_cost=part_cost,
                        personal_cost=float(hourly_rate * hours_worked),
                        cost_deduction=cost_deduction)
        end_time = time.time()

        self.__logger.info(f"OCR-Pipeline executed successfully in {round(end_time-start_time, 4)} sec.")
        return None


    #----------Private Methods----------#
    def __fill_variables(self, doc: Document, email: str = None) -> None:
        """
        Fills variables like email, invoice_nr, etc. on the invoice
        
        Parameters
        ----------
        doc : Document (docx module)
            Document where images should be written on
        email : str (default=None)
            Email adress for paypal

        Returns
        -------
        None
        """

        data = {
            "[Date]": date.today().strftime("%d-%m-%Y"),
            "[Invoice-Number]": str(random.randint(111, 1111111)),
            "[PayPal]": email
        }

        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        return None



if __name__ == "__main__":
    invoice = Invoice()
